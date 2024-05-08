from flask import Blueprint
from flask import jsonify
from flask import request
from db_utils import get_conn
import os

import docx
from flask import request



teacher_bp = Blueprint("teacher", __name__)

#将word文档读取的内容存入数据库
def read_doc(filename,form):
    con = get_conn()
    cur = con.cursor()
    try:
        doc = docx.Document(filename)
    except:
        print(f"delete from exam_manage where examCode = '{form.get('examCode')}'")
        cur.execute(f"delete from exam_manage where examCode = '{form.get('examCode')}'")
        cur.commit()
        return 'error'
    for i in range(0, len(doc.paragraphs), 7):#0-文件总行数，间隔为7(一题七行)
        sql = f"INSERT INTO multi_question (subject,examCode,question, answerA, answerB, answerC, answerD,rightAnswer) VALUES ('{form.get('source')}','{form.get('examCode')}','%s', '%s', '%s', '%s', '%s', '%s')" % \
              (doc.paragraphs[i].text, doc.paragraphs[i + 1].text, doc.paragraphs[i + 2].text,
               doc.paragraphs[i + 3].text, doc.paragraphs[i + 4].text,doc.paragraphs[i + 5].text)
        print(sql)
        try:
            cur.execute(sql)
            con.commit()
        except:
            con.rollback()
            print(form.get('examCode'))
            cur.execute(f"delete from exam_manage where examCode = '{form.get('examCode')}'")
            return 'error'
    return 'ok'

# 4、将word文档读取的内容存入数据库
def read_doc(filename,form):
    con = get_conn()
    cur = con.cursor()
    try:
        doc = docx.Document(filename)
    except:
        print(f"delete from exam_manage where examCode = '{form.get('examCode')}'")
        cur.execute(f"delete from exam_manage where examCode = '{form.get('examCode')}'")
        con.commit()
        return 'error'
    flagType = 0
    i = 0
    g = 0
    while i<len(doc.paragraphs):#0-文件总行数，间隔为7(一题8行)
        sql = ""
        if doc.paragraphs[i].text == "选择题":
            tableName = "multi_question"
            flagType = 1
            i = i+1
            g = 8
        elif doc.paragraphs[i].text == "填空题":
            tableName = "fill_question"
            flagType = 2
            i = i+1
            g = 4
        elif doc.paragraphs[i].text == "判断题":
            tableName = "judge_question"
            flagType = 3
            i = i+1
            g = 4       
        print(doc.paragraphs[i].text) 
        if flagType == 1:
            sql = f"INSERT INTO multi_question (questionType,subject,question, answerA, answerB, answerC, answerD,rightAnswer,score) VALUES (1,'{form.get('source')}','%s', '%s', '%s', '%s', '%s', '%s','%s')" % \
              (doc.paragraphs[i].text, doc.paragraphs[i + 1].text, doc.paragraphs[i + 2].text,
               doc.paragraphs[i + 3].text, doc.paragraphs[i + 4].text,doc.paragraphs[i + 5].text,doc.paragraphs[i + 6].text)
        elif flagType == 2:
            sql = f"INSERT INTO fill_question (questionType,subject,question,answer,score) VALUES (2,'{form.get('source')}','%s', '%s','%s')" % \
            (doc.paragraphs[i].text, doc.paragraphs[i + 1].text, doc.paragraphs[i + 2].text)
        elif flagType == 3:
            sql = f"INSERT INTO judge_question (questionType,subject,question, answer,score) VALUES (3,'{form.get('source')}','%s', '%s', '%s')" % \
              (doc.paragraphs[i].text, doc.paragraphs[i + 1].text, doc.paragraphs[i + 2].text)
        print(sql)
        try:
            cur.execute(sql)
            con.commit()
            paperId = getInfoFromDataBase("exam_manage","paperId","examCode",form.get('examCode'),1)['paperId']
            questionType = flagType
            questionId = getInfoFromDataBase(tableName,"questionId","",0,2)['MAX(questionId)']
            print(paperId)
            print(questionType)
            print(questionId)
            if addPaperManage(paperId,questionType,questionId) == 'error':
                return 'error'
        except:
            con.rollback()
            print(form.get('examCode'))
            # cur.execute(f"delete from exam_manage where examCode = '{form.get('examCode')}'")
            return 'error'
        i += g
    return 'ok'


#增加试卷
def addPaperManage(paperId,questionType,questionId):
    con = get_conn()
    cur = con.cursor()
    try:
        cur.execute(f"insert into paper_manage values('{paperId}','{questionType}','{questionId}')")
        con.commit()
    except:
        con.rollback()
        return 'error'
    return 'ok'

# 读取前端提交的文件
@teacher_bp.route('/upload', methods=['GET', 'POST'])
def upload():
        # file为上传表单的name属性值
        f = request.files['file'];
        data = request.form
        f.save(os.path.join(os.getcwd()+"\\src\\paper", f.filename))
        # #将题目存入数据库中
        return read_doc(os.getcwd()+"\\src\\paper\\"+f.filename,data)



#获取考试编号
@teacher_bp.route('/GetExamCode', methods=['GET', 'POST'])
def getExamManagePaperId():
    #获取数据库信息
    con = get_conn()
    cur = con.cursor()

    cur.execute("SELECT MAX(examCode) from exam_manage")
    MaxExamCode = cur.fetchone() 
    print(MaxExamCode['MAX(examCode)'])
    return MaxExamCode


#添加考试
@teacher_bp.route('/exam', methods=['GET', 'POST'])
def addExam():
    data = request.json
    source = data.get('source')
    major = data.get('major')
    grade = data.get('grade')
    examDate = data.get('examDate')
    totalTime = data.get('totalTime')
    totalScore = data.get('totalScore')
    type = data.get('type')
    examCode = data.get('examCode')
    paperId = getInfoFromDataBase("exam_manage","paperId","",0,2)
    paperId = int(paperId['MAX(paperId)'])+1
    con = get_conn()
    cur = con.cursor()
    try:
        cur.execute(f"INSERT INTO exam_manage(examCode,source,paperId,examDate,totalTime,grade,major,totalScore,type) VALUES ('{examCode}','{source}','{paperId}','{examDate}','{totalTime}','{grade}','{major}','{totalScore}','{type}')")
        con.commit()
    except:
        con.rollback()
        return '400'
    
    return '200'
 

def getInfoFromDataBase(tableName,columnName,column,attribute,Type):
    con = get_conn()
    cur = con.cursor()
    sql = ""
    if Type == 1:
        sql = f"select {columnName} from {tableName} where {column} = {attribute}"
    else:
        sql = f"SELECT MAX({columnName}) from {tableName}"
    print(sql)
    try:
        cur.execute(sql)
        con.commit()
        Info = cur.fetchone() 
    except:
        con.rollback()
    return Info


#搜索考试信息
@teacher_bp.route('/teacher/exams/<int:current>/<int:size>/<string:inputValue>', methods=['GET'])
def searchExams(current, size, inputValue):
    try:
        con = get_conn()
        cur = con.cursor()

        source = inputValue
        # 根据分页参数查询相应的考试信息
        cur.execute("SELECT * FROM exam_manage where source = %s", (source,))
        exams = cur.fetchall()

        # 获取搜索到的考试总数
        total = len(exams)

        response_data = {
            'pagination': {
                'current': current,
                'total': total,
                'size': size,
                'records': exams
            }
        }

        return jsonify({'code': 200, 'data': response_data})
    except Exception as e:
        print(e)
        return jsonify({'error': 'Internal Server Error'}), 500
    finally:
        if con:
            con.close()



#获取当前用户的考试总数
@teacher_bp.route('/teacher/exams/<int:current>/<int:size>', methods=['GET'])
def get_exams(current, size):
    try:
        print("进入")
        con = get_conn()
        cur = con.cursor()

        # 根据分页参数查询相应的考试信息
        cur.execute("SELECT * FROM exam_manage LIMIT %s OFFSET %s", (size, (current - 1) * size))
        exams = cur.fetchall()
        # print(exams)

        # 获取考试总数
        cur.execute("SELECT COUNT(*) as total FROM exam_manage")
        total = cur.fetchone()['total']

        response_data = {
            'pagination':{
                'current':current,
                'total':total,
                'size':size,
                'records': exams
            }
        }

        return jsonify({'code': 200,'data': response_data})
    except Exception as e:
        print(e)
        return jsonify({'error': 'Internal Server Error'}), 500
    finally:
        if con:
            con.close()

# 获取考试详情页
@teacher_bp.route('/teacher/exam/<int:examCode>', methods=['GET'])
def get_examDetails(examCode):
    try:
        print(examCode)
        con = get_conn()
        cur = con.cursor()

        # 查询试卷详细信息
        cur.execute("SELECT * FROM exam_manage WHERE examCode = %s", (examCode,))
        exam_details = cur.fetchone()
        print(exam_details)

        # 查询试卷的题目信息（这里假设有个 paperId 字段关联试卷和题目，你需要根据实际情况调整）
        paperId = exam_details.get('paperId')
        cur.execute("SELECT * FROM paper_manage WHERE paperId = %s", (paperId,))
        paper_details = cur.fetchall()

        response_data = {
            'examData': exam_details,
            'paperId': paperId
        }

        return jsonify({'code':200,'data': response_data})
    except Exception as e:
        print(e)
        return jsonify({'error': 'Internal Server Error'}), 500
    finally:
        if con:
            con.close()

#删除考试
@teacher_bp.route('/teacher/exam/delete/<int:examCode>', methods=['POST'])
def delete_exam(examCode):
    try:
        con = get_conn()
        cur = con.cursor()

        # 查询试卷详细信息
        cur.execute("DELETE FROM exam_manage WHERE examCode = %s", (examCode,))
        con.commit()
        return jsonify({'code':200})
    except Exception as e:
        print(e)
        return jsonify({'error': 'Internal Server Error'}), 500
    finally:
        if con:
            con.close()

# 编辑考试信息
@teacher_bp.route('/teacher/exam', methods=['POST','GET'])
def edit_examDetails():
    try:
        con = get_conn()
        cur = con.cursor()

        teacher_data = request.get_json()
        examCode = teacher_data.get('examCode')
        cur.execute("""
            UPDATE exam_manage
            SET
                source = %s,
                description = %s,
                institute = %s,
                major = %s,
                grade = %s,
                examDate = %s,
                totalTime = %s,
                totalScore = %s,
                `type` = %s,
                tips = %s
            WHERE examCode = %s
        """, (
            teacher_data.get('source'),
            teacher_data.get('description'),
            teacher_data.get('institute'),
            teacher_data.get('major'),
            teacher_data.get('grade'),
            teacher_data.get('examDate'),
            teacher_data.get('totalTime'),
            teacher_data.get('totalScore'),
            teacher_data.get('type'),
            teacher_data.get('tips'),
            examCode
        ))

        # Commit changes to the database
        con.commit()
        return jsonify({'code': 200})
    except Exception as e:
        print(e)
        return jsonify({'error': 'Internal Server Error'}), 500
    finally:
        if con:
            con.close()

#添加教师
@teacher_bp.route('/teacher/add', methods=['POST'])
def add_teacher():
    try:
        con = get_conn()
        cur = con.cursor()

        # 从前端请求中获取数据
        teacher_data = request.get_json()

        # 将数据插入到数据库中
        cur.execute("""
            INSERT INTO teacher (teacherName, institute, tel, email, pwd, cardId, sex, type, role)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
        """, (
            teacher_data['teacherName'],
            teacher_data['institute'],
            teacher_data['tel'],
            teacher_data['email'],
            teacher_data['pwd'],
            teacher_data['cardId'],
            teacher_data['sex'],
            teacher_data['type'],
            teacher_data['role']
        ))

        con.commit()

        return jsonify({'code': 200})
    except Exception as e:
        print(e)
        return jsonify({'error': 'Internal Server Error'}), 500
    finally:
        if con:
            con.close()

#获取教师信息
@teacher_bp.route('/seacherteacher/<int:teacherId>', methods=['GET'])
def get_teacherDetails(teacherId):
    try:
        con = get_conn()
        cur = con.cursor()

        # 根据分页参数查询相应的教师信息
        cur.execute("SELECT * FROM teacher where teacherId = %s", teacherId)
        teacher_details = cur.fetchone()
        print(teacher_details)


        response_data = {
            'records':teacher_details
        }


        return jsonify({'code': 200,'data': response_data})
    except Exception as e:
        print(e)
        return jsonify({'error': 'Internal Server Error'}), 500
    finally:
        if con:
            con.close()

#获取教师总数
@teacher_bp.route('/teachers/<int:current>/<int:size>', methods=['GET'])
def get_teachers(current, size):
    try:
        print("进入")
        con = get_conn()
        cur = con.cursor()

        # 根据分页参数查询相应的教师信息
        cur.execute("SELECT * FROM teacher LIMIT %s OFFSET %s", (size, (current - 1) * size))
        teachers = cur.fetchall()
        # print(exams)

        # 获取教师总数
        cur.execute("SELECT COUNT(*) as total FROM teacher")
        total = cur.fetchone()['total']

        print(total)

        response_data = {
            'pagination':{
                'current':current,
                'total':total,
                'size':size,
                'records': teachers
            }
        }


        return jsonify({'code': 200,'data': response_data})
    except Exception as e:
        print(e)
        return jsonify({'error': 'Internal Server Error'}), 500
    finally:
        if con:
            con.close()

#删除教师
@teacher_bp.route('/teacher/delete/<int:teacherId>', methods=['DELETE'])
def delete_teacher(teacherId):
    try:
        con = get_conn()
        cur = con.cursor()

        # 查询试卷详细信息
        cur.execute("DELETE FROM teacher WHERE teacherId = %s", (teacherId,))
        con.commit()
        return jsonify({'code':200})
    except Exception as e:
        print(e)
        return jsonify({'error': 'Internal Server Error'}), 500
    finally:
        if con:
            con.close()

# 编辑教师信息
@teacher_bp.route('/teacher/edit', methods=['POST'])
def edit_teacherDetails():
    try:
        con = get_conn()
        cur = con.cursor()

        # 从前端请求中获取数据
        teacher_data = request.get_json()
        teacherId = request.get_json().get('teacherId')

        cur.execute("""
            UPDATE teacher
            SET
                teacherName = %s,
                institute = %s,
                sex = %s,
                tel = %s,
                email = %s,
                pwd = %s,
                cardId = %s,
                `type` = %s
            WHERE teacherId = %s
        """, (
            teacher_data.get('teacherName'),
            teacher_data.get('institute'),
            teacher_data.get('sex'),
            teacher_data.get('tel'),
            teacher_data.get('email'),
            teacher_data.get('pwd'),
            teacher_data.get('cardId'),
            teacher_data.get('type'),
            teacherId
        ))

        # Commit changes to the database
        con.commit()
        return jsonify({'code': 200})
    except Exception as e:
        print(e)
        return jsonify({'error': 'Internal Server Error'}), 500
    finally:
        if con:
            con.close()

#
@teacher_bp.route('/student/<int:current>/<int:size>', methods=['GET', 'POST'])
def getAllStudentInfo(current,size):
    con = get_conn()
    cur = con.cursor()
    try:
        cur.execute(f"SELECT * FROM student LIMIT {(current-1)*size},{size}")
        con.commit()
        studentInfo = cur.fetchall()
    except:
        con.rollback()
        return 'error'
    return studentInfo

@teacher_bp.route('/student/<int:studentId>', methods=['DELETE'])
def deleteStudentInfo(studentId):
    con = get_conn()
    cur = con.cursor()
    try:
        cur.execute(f"DELETE FROM student WHERE studentId = {studentId}")
        con.commit()
    except:
        con.rollback()
        return 'error'
    return 'ok'

@teacher_bp.route('/student', methods=['PUT'])
def modifiedStudentInfo():
    data = request.json
    studentId = data.get('studentId')
    studentName = data.get('studentName')
    institute = data.get('institute')
    major = data.get('major')
    grade = data.get('grade')
    clazz = data.get('clazz')
    sex = data.get('sex')
    tel = data.get('tel')
    con = get_conn()
    cur = con.cursor()
    try:
        cur.execute(f"UPDATE student set studentName='{studentName}',institute='{institute}',major='{major}',grade='{grade}',clazz='{clazz}',sex='{sex}',tel='{tel}' where studentId = '{studentId}'")
        con.commit()
    except:
        con.rollback()
        return '400'
    return '200'

@teacher_bp.route('/student/<int:studentId>', methods=['GET', 'POST'])
def getStudentInfo(studentId):
    con = get_conn()
    cur = con.cursor()
    try:
        cur.execute(f"SELECT * FROM student WHERE studentId = {studentId}")
        con.commit()
        studentInfo = cur.fetchone()
    except:
        con.rollback()
        return 'error'
    return studentInfo

@teacher_bp.route('/student', methods=['POST'])
def addStudent():
    data = request.json
    studentName = data.get('studentName')
    institute = data.get('institute')
    major = data.get('major')
    grade = data.get('grade')
    clazz = data.get('clazz')
    sex = data.get('sex')
    tel = data.get('tel')
    email = data.get('email')
    pwd = data.get('pwd')
    cardId = data.get('cardId')
    role = data.get('role')
    con = get_conn()
    cur = con.cursor()
    try:
        cur.execute("INSERT INTO student(studentName,grade,major,clazz,institute,tel,email,pwd,cardId,sex,role)"+
                     f"VALUES('{studentName}','{grade}','{major}','{clazz}','{institute}','{tel}','{email}','{pwd}','{cardId}','{sex}','{role}')")
        con.commit()
    except:
        con.rollback()
        print("INSERT INTO student(studentName,grade,major,clazz,institute,tel,email,pwd,cardId,sex,role)"+
                     f"VALUES('{studentName}','{grade}','{major}','{clazz}','{institute}','{tel}','{email}','{pwd}','{cardId}','{sex}','{role}'")
        return '400'
    return '200'

#根据查询条件查询学生信息
@teacher_bp.route('/condition', methods=['GET', 'POST'])
def getSpecificStudentInfo():
    data = request.json
    current = data.get('current')
    size = data.get('size')
    student = data.get('inputCondition')
    major = data.get('major')
    grade = data.get('grade')
    clazz = data.get('clazz')
    con = get_conn()
    cur = con.cursor()
    try:
        frontSql = "SELECT * FROM student WHERE 1=1 "
        studentCondition = ""
        majorCondition = ""
        gradeCondition = ""
        clazzCondition = ""
        if not(str(student).isspace() or str(student) == ""):
            studentCondition = f"AND (studentId LIKE '%{student}%' OR studentName LIKE '%{student}%') "
        if not(str(major) == "无" or str(major) == ""):
            majorCondition = f"AND (major = '{major}') "
        if not(str(grade) == "无" or str(grade) == ""):
            gradeCondition = f"AND (grade = '{grade}') " 
        if not(str(clazz) == "无" or str(clazz) == ""):
            clazzCondition = f"AND (clazz = '{clazz}') "         
        limit = f"LIMIT {size} OFFSET {(current-1)*(size)}"
        sql = frontSql + studentCondition + majorCondition + gradeCondition + clazzCondition + limit
        print(sql)
        cur.execute(sql)
        con.commit()
        studentInfo = cur.fetchall()
    except:
        con.rollback()
        return jsonify({'code':400})
    return jsonify({'code':200 ,'data':studentInfo})


#根据学生ID查询该学生的所有成绩
@teacher_bp.route('/score/<int:studentId>', methods=['GET'])
def get_scoreDetails(studentId):
    print(studentId)
    try:
        #与数据库建立连接
        con = get_conn()
        cur = con.cursor()

         # 执行查询学生所有考试名称及对应的考试成绩
        query = "SELECT exam_manage.source, score.examCode, score.etScore FROM score INNER JOIN exam_manage ON score.examCode = exam_manage.examCode WHERE score.studentId = %s"
        cur.execute(query, (studentId,))
        rows = cur.fetchall()
        print(rows)

        # 将查询结果整理成前端需要的格式
        score_data = [{'exam_name':row['source'],'exam_number': row['examCode'], 'exam_score': row['etScore']} for row in rows]
        print(score_data)

        return jsonify({'code': 200, 'data': score_data})

    except Exception as e:
        print(str(e))
        return jsonify({'code': 500, 'message': 'Error occurred while fetching scores'})
    finally:
        cur.close()
        con.close()

#调出所有考试信息
@teacher_bp.route('/teachers/exams/<int:current>/<int:size>', methods=['GET'])
def get_all_exams(current, size):
    print (current)
    print (size)
    print("查询考试信息")
    try:
        con = get_conn()
        cur = con.cursor()

        # 根据当前页和每页大小计算偏移量
        offset = (current - 1) * size

        # 从数据库中检索考试的相关信息
        query = "SELECT source, examCode, description, institute, major, grade, examDate, totalTime,totalScore FROM exam_manage LIMIT %s OFFSET %s"
        cur.execute(query, (size, offset))
        rows = cur.fetchall()

        # 检索用于分页的记录总数
        total_query = "SELECT COUNT(*) FROM exam_manage"
        cur.execute(total_query)
        total = cur.fetchone()['COUNT(*)']

        # 将结果以字典格式准备好
        result = {
            'current': current,
            'total': total,
            'size': size,
            'data': [
                {
                    'source': row['source'],
                    'examCode': row['examCode'],
                    'description': row['description'],
                    'institute': row['institute'],
                    'major': row['major'],
                    'grade': row['grade'],
                    'examDate': row['examDate'],
                    'totalTime': row['totalTime'],
                    'totalScore': row['totalScore'],
                } for row in rows
            ]
        }
        return jsonify({'code': 200, 'data': result})
    except Exception as e:
        print(e)
        return jsonify({'error': '内部服务器错误'}), 500
    finally:
        cur.close()
        con.close()

#根据考试编号查询考试情况
@teacher_bp.route('/teacher/scores/<int:examCode>', methods=['GET'])
def get_scoreParts(examCode):
    print(examCode)
    try:
        con = get_conn()
        cur = con.cursor()

        # 查询数据库，获取对应考试编号的所有成绩信息
        query = "SELECT etScore FROM score WHERE examCode = %s"
        cur.execute(query,(examCode))
        scores = cur.fetchall()
        print(scores)

        # 返回 JSON 格式的数据给前端
        return jsonify({'code': 200, 'data': scores, 'message': 'Success'})

    except Exception as e:
        print(str(e))
        return jsonify({'code': 500, 'message': 'Error occurred while fetching scores'})
    finally:
        cur.close()
        con.close()

#根据查询条件查询学生成绩数据
@teacher_bp.route('/Rankcondition/<int:examCode>', methods=['GET', 'POST'])
def getSpecificStudentInfos(examCode):
    print("成功进入")
    data = request.json
    current = data.get('current')
    size = data.get('size')
    student = data.get('inputCondition')
    major = data.get('major')
    grade = data.get('grade')
    clazz = data.get('clazz')
    con = get_conn()
    cur = con.cursor()
    
    try:
        frontSql = """
            SELECT 
                student.studentId, student.studentName, student.major, 
                student.grade, student.clazz, student.sex, score.score
            FROM 
                student
            JOIN 
                score ON student.studentId = score.studentId
            WHERE 
                1=1 AND score.examCode = {}
        """.format(examCode)
        studentCondition = ""
        majorCondition = ""
        gradeCondition = ""
        clazzCondition = ""
        
        if not(str(student).isspace() or str(student) == ""):
            studentCondition = f"AND (studentId LIKE '%{student}%' OR studentName LIKE '%{student}%') "
        if not(str(major) == "无" or str(major) == ""):
            majorCondition = f"AND (major = '{major}') "
        if not(str(grade) == "无" or str(grade) == ""):
            gradeCondition = f"AND (grade = '{grade}') " 
        if not(str(clazz) == "无" or str(clazz) == ""):
            clazzCondition = f"AND (clazz = '{clazz}') "    
        limit = f"LIMIT {size} OFFSET {(current-1)*(size)}"
        sql = frontSql + studentCondition + majorCondition + gradeCondition + clazzCondition + limit
        print(sql)
        cur.execute(sql)
        con.commit()
        studentInfo = cur.fetchall()
        print(studentInfo)
    except:
        con.rollback()
        return jsonify({'code': 400})
    return jsonify({'code': 200, 'data': studentInfo})

#按专业查询学生信息
@teacher_bp.route('/major', methods=['GET', 'POST'])
def getMajorInfo():
    con = get_conn()
    cur = con.cursor()
    try:
        cur.execute(f"SELECT DISTINCT major FROM student")
        con.commit()
        studentInfo = cur.fetchall()
    except:
        con.rollback()
        return jsonify({'code':400})
    return jsonify({'code':200 ,'data':studentInfo})

#按年级查询学生信息
@teacher_bp.route('/grade', methods=['GET', 'POST'])
def getGradeInfo():
    con = get_conn()
    cur = con.cursor()
    try:
        cur.execute(f"SELECT DISTINCT grade FROM student")
        con.commit()
        studentInfo = cur.fetchall()
    except:
        con.rollback()
        return jsonify({'code':400})
    return jsonify({'code':200 ,'data':studentInfo})

#按班级查询学生信息
@teacher_bp.route('/clazz', methods=['GET', 'POST'])
def getClazzInfo():
    con = get_conn()
    cur = con.cursor()
    try:
        cur.execute(f"SELECT DISTINCT clazz FROM student")
        con.commit()
        studentInfo = cur.fetchall()
    except:
        con.rollback()
        return jsonify({'code':400})
    return jsonify({'code':200 ,'data':studentInfo})

